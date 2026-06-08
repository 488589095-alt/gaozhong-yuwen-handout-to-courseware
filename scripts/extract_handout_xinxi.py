# -*- coding: utf-8 -*-
"""信息类讲义(第10讲)提取器：docx → content.json。
信息类 vs 古诗：诗→长材料(material，渲染时分页)；题型有 选择/补写/简答；考情4列；有教材链接。
练习结构规整可自动提取；例题结构特殊(必做1+选做(1)(2))，本脚本也尽力提取，产出后人工核对。
"""
import json, re, sys
from docx import Document
from docx.oxml.ns import qn

DOCX = sys.argv[1] if len(sys.argv) > 1 else \
    "/Users/gaotu/Desktop/高中讲义&课件-语文/讲义/第10讲【信息】强化信息提取，破解情境填空 (1).docx"
OUT = sys.argv[2] if len(sys.argv) > 2 else "content.json"

d = Document(DOCX)
paras = [p.text.strip() for p in d.paragraphs]
ANS, ANA = "【答案】", "【解析】"
QRE = re.compile(r'^[（(]\s*\d+\s*[)）]')
NOTE = re.compile(r'【[^】]{0,40}】')   # 材料里的编者批注


def tbl(i):
    t = d.tables[i]
    rows = [[c.paragraphs[0].text.strip() for c in r.cells] for r in t.rows]
    return rows[0], rows[1:]


def section(start, end):
    si = paras.index(start)
    out = []
    for t in paras[si + 1:]:
        if t == end or t.startswith(end[:4]):
            break
        if t:
            out.append(t)
    return out


def is_choice(qtext):
    return ("一项是" in qtext) or ("正确的一项" in qtext)


def parse_block(lines):
    """一个练习/例题块的 lines → (material, questions[])。material 截到首个 (n)。"""
    qi = next((i for i, t in enumerate(lines) if QRE.match(t)), len(lines))
    material = "\n".join(NOTE.sub("", t).strip() for t in lines[:qi]
                         if NOTE.sub("", t).strip())
    # 题干：(n) 起，累积到下一个 (n) 或 【答案/解析】
    qs, cur = [], None
    ans = []
    in_ans = False
    collecting = False          # 是否在并入"当前简答答案"的续行
    STOP = ("【解析", "【考", "【试", "【其他", "本题考查", "故选", "考试院", "解析二", "解析（")
    for t in lines[qi:]:
        if t.startswith(ANS):
            body = t[len(ANS):].strip()
            in_ans = True
            if body:
                ans.append(body)
                collecting = not re.fullmatch(r'[A-D]', body)   # 字母答案不再续行
            else:
                collecting = False
            continue
        if t.startswith(STOP):      # 进入解析，停止续行
            in_ans = True
            collecting = False
            continue
        if in_ans:
            if collecting and ans:
                ans[-1] += "\n" + t
            continue
        if QRE.match(t):
            if cur:
                qs.append(cur)
            cur = t
        elif cur is not None:
            cur += "\n" + t
    if cur:
        qs.append(cur)
    # 配对答案（按序），判类型
    out = []
    for i, qt in enumerate(qs):
        a = ans[i] if i < len(ans) else ""
        if is_choice(qt) and re.fullmatch(r'[A-D]', a.strip()):
            out.append({"type": "choice", "text": qt, "answer_letter": a.strip()})
        else:
            out.append({"type": "essay", "text": qt,
                        "answer": (ANS + a) if a else ANS})
    return material, out


# ── 表格 ──
oh, orow = tbl(0)
eh, erow = tbl(1)

# ── 教材链接 ──
tb_paras = section("学习内容", "知识方法")
# 去掉"知识方法"下属的(一)(二)，只要学习内容正文段
textbook = [p for p in tb_paras if not p.startswith(("（一）", "（二）", "知识方法"))]

# ── 针对练习 ──
pi = paras.index("针对练习")
prac_lines = paras[pi + 1:]
# 按 "练习N阅读" 切分
idxs = [i for i, t in enumerate(prac_lines) if re.match(r'^练习\s*\d+\s*阅读', t)]
practices = []
labels = {0: "【课前测】", 3: None}   # 标签按需补
for k, st in enumerate(idxs):
    en = idxs[k + 1] if k + 1 < len(idxs) else len(prac_lines)
    block = prac_lines[st:en]
    head = block[0]
    src = "【真题】阅读下面的文字，完成下面的小题。"
    material, qs = parse_block(block[1:])
    practices.append({"index": f"{k+1:02d}", "source": src,
                      "material": material, "questions": qs})

# ── 例题（结构特殊：必做补写 + 选做(1)选择 +(2)简答）──
ei = paras.index("例题分析")
ex_lines = paras[ei + 1:pi]          # pi=针对练习
ex_src = ex_lines[0]
one_idx = next(i for i, t in enumerate(ex_lines) if t.strip() == "1")
ex_material = "\n".join(NOTE.sub("", t).strip() for t in ex_lines[1:one_idx]
                        if NOTE.sub("", t).strip())
rest = ex_lines[one_idx + 1:]
# 例题答案（同 collecting 逻辑），顺序：必做 / 选做1字母 / 选做2
ex_ans, collecting = [], False
STOP2 = ("【解析", "【考", "【试", "【其他", "本题考查", "故选", "考试院")
for t in rest:
    if t.startswith(ANS):
        b = t[len(ANS):].strip()
        if b:
            ex_ans.append(b); collecting = not re.fullmatch(r'[A-D]', b)
        else:
            collecting = False
    elif t.startswith(STOP2):
        collecting = False
    elif collecting and ex_ans:
        ex_ans[-1] += "\n" + t
must_stem = next((t for t in rest if "补写出恰当" in t), "")
must_seg = rest[rest.index(must_stem) + 1] if must_stem in rest else ""
ci = next(i for i, t in enumerate(rest) if "一项是" in t)
cblock = [rest[ci]]
for t in rest[ci + 1:]:
    if re.match(r'^[A-D][．.]', t):
        cblock.append(t)
    else:
        break
essay_stem = next((t for t in rest if "清代的古籍" in t), "")
example = {"source": ex_src, "material": ex_material, "questions": [
    {"tag": "【必做题】", "type": "essay", "text": must_stem + "\n" + must_seg,
     "answer": ANS + (ex_ans[0] if len(ex_ans) > 0 else "")},
    {"tag": "【选做题】", "type": "choice", "text": "\n".join(cblock),
     "answer_letter": ex_ans[1] if len(ex_ans) > 1 else ""},
    {"tag": "【选做题】", "type": "essay", "text": essay_stem,
     "answer": ANS + (ex_ans[2] if len(ex_ans) > 2 else "")},
]}

content = {
    "lecture_no": "第10讲",
    "title": "【信息】强化信息提取，破解情境填空",
    "grade": "高一", "teacher": "主讲老师：", "term": "2026寒春",
    "system_name": "学习成长与规划系统",
    "modules": ["学习目标", "考情分析", "教材链接", "典型例题", "针对练习"],
    "objectives": {"header": oh, "rows": orow},
    "star_legend": ["★★★★★  120+分必会", "★★★★✩  90-120分必会",
                    "★★★✩✩  60-90分必会", "★★✩✩✩  60分必会"],
    "exam": {"header": eh, "rows": erow},
    "textbook_link": {"title": "教材链接", "paragraphs": textbook},
    "example": example,
    "practices": practices,
    "knowledge": {"steps": ["1. 强化信息提取能力", "2. 破解情境填空题型"]},
    "end": {"big": "下节课我们\n再见啦～", "small": "本期课结束"},
}
json.dump(content, open(OUT, "w", encoding="utf-8"), ensure_ascii=False, indent=2)
print(f"✅ {OUT}")
print(f"  学习目标 {len(orow)}行 | 考情 {len(erow)}行({len(eh)}列) | 教材链接 {len(textbook)}段")
print(f"  练习 {len(practices)} 个：")
for p in practices:
    print(f"   练习{p['index']}: 材料{len(p['material'])}字, {len(p['questions'])}问 "
          f"[{','.join(q['type'][0] for q in p['questions'])}] "
          f"答案letter={[q.get('answer_letter') for q in p['questions'] if q['type']=='choice']}")
