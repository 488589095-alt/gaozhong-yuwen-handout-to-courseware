# -*- coding: utf-8 -*-
"""
extract_handout.py —— 从高一语文讲义 docx 自动抽取结构化内容（朝通用 skill）

抽取：
  - knowledge：炼字概述 / 常见命题角度+设问方式 / 常考词类作用表(7类) / 解题步骤4步
  - qa：按题目出现顺序的 [{answer, analysis}]（取讲义里非空的【答案】【解析】）

输出 handout_extracted.json。可加 --merge 合并进 content.json
（knowledge 整块 + 按顺序把 analysis 写入每题）。
"""
import argparse, json, re
from pathlib import Path
from docx import Document

ANS, ANA = "【答案】", "【解析】"


def _texts(doc):
    return [p.text.strip() for p in doc.paragraphs]


def extract_knowledge(doc, paras):
    def between(a, b):
        try:
            i = paras.index(a)
        except ValueError:
            return []
        out = []
        for t in paras[i + 1:]:
            if t == b or t.startswith(b[:4]):
                break
            if t:
                out.append(t)
        return out

    concept = " ".join(between("一、炼字概述", "二、常见命题角度"))
    ang = between("二、常见命题角度", "三、常考词类作用")
    angle_intro = ang[0] if ang else ""
    ask_ways = [a for a in ang[1:] if a and not a.endswith("：")]  # 去掉"其一般设问方式是："

    # 词类作用表 = 文档第3个表格（表0学习目标 表1考情 表2词类）
    wt = None
    for t in doc.tables:
        hdr = [c.paragraphs[0].text.strip() for c in t.rows[0].cells]
        if "词" in hdr[0] and ("作用" in "".join(hdr) or len(t.rows) >= 7):
            wt = t
            break
    word_types = None
    if wt is not None:
        rows = []
        for r in wt.rows:
            rows.append([c.paragraphs[0].text.strip() for c in r.cells])
        word_types = {"header": rows[0], "rows": rows[1:]}
    return {"concept": concept, "angle_intro": angle_intro,
            "ask_ways": ask_ways, "word_types": word_types}


def extract_qa(paras):
    """状态机：收集非空【答案】【解析】，遇题块边界停止累积。"""
    try:
        start = paras.index("例题分析")
    except ValueError:
        start = 0
    qa = []
    cur_ans = cur_ana = None
    mode = None

    def is_boundary(t):
        return (re.match(r'^练习\s*\d', t) or re.match(r'^（\s*\d\s*）', t)
                or re.match(r'^\(\s*\d\s*\)', t) or t in ("针对练习", "例题分析", "1", "2", "3")
                or re.match(r'^[A-DＡ-Ｄ][．.、]', t))

    def flush():
        nonlocal cur_ans, cur_ana
        if cur_ans:
            qa.append({"answer": cur_ans.strip(), "analysis": (cur_ana or "").strip()})
        cur_ans = cur_ana = None

    for t in paras[start + 1:]:
        if t.startswith(ANS):
            body = t[len(ANS):].strip()
            if not body:        # 空占位
                continue
            flush()
            cur_ans, cur_ana, mode = body, None, "ans"
        elif t.startswith(ANA):
            body = t[len(ANA):].strip()
            if not body:
                continue
            cur_ana, mode = body, "ana"
        elif not t:
            continue
        elif is_boundary(t):
            mode = None          # 进入新题区，停止累积（等下个【答案】）
        else:
            if mode == "ans" and cur_ans is not None:
                cur_ans += "\n" + t
            elif mode == "ana" and cur_ana is not None:
                cur_ana += "\n" + t
    flush()
    return qa


def steps_from(qa):
    """解题步骤4步标题：从例1解析里抓"第X步，YYY，ZZZ。"。"""
    steps = []
    src = qa[0]["analysis"] if qa else ""
    for m in re.finditer(r'第([一二三四五六])步[，,]\s*([^。\n]+)。', src):
        steps.append(f"第{m.group(1)}步  {m.group(2)}")
    return steps


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True)
    ap.add_argument("-o", "--out", default="handout_extracted.json")
    ap.add_argument("--merge", help="合并进该 content.json")
    a = ap.parse_args()

    doc = Document(a.docx)
    paras = _texts(doc)
    knowledge = extract_knowledge(doc, paras)
    qa = extract_qa(paras)
    knowledge["steps"] = steps_from(qa)

    out = {"knowledge": knowledge, "qa": qa}
    Path(a.out).write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"✅ 抽取完成 → {a.out}")
    print(f"  概述: {knowledge['concept'][:30]}…")
    print(f"  设问方式 {len(knowledge['ask_ways'])} 条")
    wt = knowledge["word_types"]
    print(f"  词类表: {len(wt['rows']) if wt else 0} 行  表头={wt['header'] if wt else None}")
    print(f"  解题步骤: {knowledge['steps']}")
    print(f"  qa 共 {len(qa)} 组：")
    for i, x in enumerate(qa):
        print(f"   [{i}] 答:{x['answer'][:22]!r}  解析:{x['analysis'][:18]!r}")

    if a.merge:
        merge_into_content(out, a.merge)


def merge_into_content(extracted, content_path):
    C = json.loads(Path(content_path).read_text(encoding="utf-8"))
    C["knowledge"] = extracted["knowledge"]
    # 按顺序铺 analysis：example.questions + 各 practice.questions
    qlist = list(C["example"]["questions"])
    for p in C["practices"]:
        qlist += list(p["questions"])
    qa = extracted["qa"]
    if len(qlist) != len(qa):
        print(f"⚠ 题数不匹配 content={len(qlist)} vs 抽取={len(qa)}，跳过 analysis 合并")
    else:
        for q, x in zip(qlist, qa):
            q["analysis"] = x["analysis"]
        print(f"✅ analysis 已按序合并到 {len(qa)} 题")
    Path(content_path).write_text(json.dumps(C, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"✅ knowledge + analysis 已写入 {content_path}")


if __name__ == "__main__":
    main()
